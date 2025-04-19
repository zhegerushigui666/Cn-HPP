import os
from docx import Document
from .utils import is_chinese

class FileHandler:
    """文件处理基类"""
    def __init__(self):
        self.entities = []
        
    def redact(self, input_path, output_path, strategy, language=None):
        """
        处理文件并替换隐私信息
        
        参数:
            input_path: 输入文件路径
            output_path: 输出文件路径
            strategy: 使用的识别策略
            language: 指定语言，如果为None则自动检测
        """
        raise NotImplementedError
        
    def get_entities(self):
        """获取最近一次处理中识别的实体"""
        return self.entities


class TextFileHandler(FileHandler):
    """处理纯文本文件"""
    
    def redact(self, input_path, output_path, strategy, language=None):
        """处理纯文本文件并替换隐私信息"""
        # 读取文本文件
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # 自动检测语言
        if language is None:
            language = 'zh' if is_chinese(text) else 'en'
        
        # 提取实体
        self.entities = strategy.extract_entities(text, language)
        
        # 替换文本
        redacted_text = text
        for entity in self.entities:
            redacted_text = redacted_text.replace(entity['original'], entity['replacement'])
        
        # 写入处理后的文本
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(redacted_text)
        
        print(f"✅ 成功处理文本文件: {output_path}")


class DocxFileHandler(FileHandler):
    """处理Word文档文件"""
    
    def redact(self, input_path, output_path, strategy, language=None):
        """处理Word文档文件并替换隐私信息"""
        # 读取文档
        doc = Document(input_path)
        
        # 处理所有段落和表格
        self.entities = []
        
        # 处理段落
        for para in doc.paragraphs:
            self._process_paragraph(para, strategy, language)
            
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._process_paragraph(para, strategy, language)
        
        # 保存处理后的文档
        try:
            doc.save(output_path)
            print(f"✅ 成功处理Word文档: {output_path}")
        except Exception as e:
            print(f"❌ 保存Word文档失败: {e}")
            
    def _process_paragraph(self, para, strategy, language):
        """处理单个段落"""
        # 获取段落文本
        text = para.text.strip()
        if not text:
            return
            
        # 自动检测语言
        if language is None:
            lang = 'zh' if is_chinese(text) else 'en'
        else:
            lang = language
            
        # 提取实体
        paragraph_entities = strategy.extract_entities(text, lang)
        if not paragraph_entities:
            return
            
        # 添加到总实体列表
        self.entities.extend(paragraph_entities)
        
        # 替换文本
        replaced_text = text
        for entity in paragraph_entities:
            replaced_text = replaced_text.replace(entity['original'], entity['replacement'])
            
        # 如果没有变化，不需要更新
        if replaced_text == text:
            return
            
        # 替换段落内容，保留格式
        self._replace_with_runs(para, replaced_text)
        
    def _replace_with_runs(self, para, new_text):
        """替换段落内容，尽量保留原始格式"""
        # 如果只有一个run，直接替换
        if len(para.runs) == 1:
            para.runs[0].text = new_text
            return
            
        # 否则，按照原始run的长度分配新文本
        offset = 0
        for run in para.runs:
            run_len = len(run.text)
            # 如果偏移量已经超过新文本长度，清空这个run
            if offset >= len(new_text):
                run.text = ''
                continue
                
            # 计算这个run应该分配多少新文本
            if offset + run_len <= len(new_text):
                run.text = new_text[offset:offset + run_len]
            else:
                run.text = new_text[offset:]
                offset = len(new_text)
                continue
                
            offset += run_len
            
        # 如果新文本更长，将剩余部分添加到最后一个run
        if offset < len(new_text):
            para.runs[-1].text += new_text[offset:] 